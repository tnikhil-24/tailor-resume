import os
from datetime import date
from docx import Document
from .config import RESUME_PATH

_PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))


def generate(decisions: dict, job_title: str, company: str, source_path: str = None) -> str:
    source = source_path or os.path.abspath(RESUME_PATH)
    doc = Document(source)

    summary = decisions.get("summary", {})
    if summary.get("accepted") and summary.get("suggested"):
        para = doc.paragraphs[summary["paragraph_index"]]
        runs = para.runs
        if runs:
            runs[0].text = summary["suggested"]
            for run in runs[1:]:
                run.text = ""
        else:
            para.add_run(summary["suggested"])

    for bullet in decisions.get("experience", []):
        if bullet.get("accepted") and bullet.get("suggested"):
            para = doc.paragraphs[bullet["paragraph_index"]]
            runs = para.runs
            if runs:
                runs[0].text = bullet["suggested"]
                for run in runs[1:]:
                    run.text = ""
            else:
                para.add_run(bullet["suggested"])

    skills = decisions.get("skills", {})
    reordered = skills.get("reordered_categories", [])
    if reordered:
        additions_by_category = {}
        for addition in skills.get("additions", []):
            if addition.get("added"):
                additions_by_category.setdefault(addition["category"], []).append(addition["skill"])

        # Reorder by writing each category's content into the i-th skill paragraph slot
        paragraph_indices = sorted(cat["paragraph_index"] for cat in reordered)
        for i, cat_data in enumerate(reordered):
            target_idx = paragraph_indices[i]
            para = doc.paragraphs[target_idx]
            items = list(cat_data["items"]) + additions_by_category.get(cat_data["category"], [])
            new_text = f"{cat_data['category']}: {', '.join(items)}"
            runs = para.runs
            if runs:
                runs[0].text = new_text
                for run in runs[1:]:
                    run.text = ""
            else:
                para.add_run(new_text)

    today = date.today().strftime("%Y-%m-%d")
    safe_company = company.strip().replace(" ", "_")
    safe_title = job_title.strip().replace(" ", "_")
    filename = f"Resume_{safe_company}_{safe_title}_{today}.docx"
    output_path = os.path.join(_PROJECT_ROOT, filename)

    doc.save(output_path)
    return output_path
