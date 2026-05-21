import os
import tempfile
import pytest
from docx import Document
from docx.shared import Pt
from backend.writer import generate


def make_test_docx(summary_text: str, bold: bool = False) -> str:
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run(summary_text)
    run.bold = bold
    run.font.size = Pt(11)

    path = tempfile.mktemp(suffix=".docx")
    doc.save(path)
    return path


def test_writer_applies_accepted_summary():
    source = make_test_docx("Original summary text.")
    decisions = {"summary": {"accepted": True, "suggested": "New tailored summary.", "paragraph_index": 0}}

    output_path = generate(decisions, "SWE", "Google", source_path=source)
    try:
        result = Document(output_path)
        assert result.paragraphs[0].text == "New tailored summary."
    finally:
        os.remove(source)
        os.remove(output_path)


def test_writer_skips_rejected_summary():
    source = make_test_docx("Original summary text.")
    decisions = {"summary": {"accepted": False, "suggested": "New tailored summary.", "paragraph_index": 0}}

    output_path = generate(decisions, "SWE", "Google", source_path=source)
    try:
        result = Document(output_path)
        assert result.paragraphs[0].text == "Original summary text."
    finally:
        os.remove(source)
        os.remove(output_path)


def test_writer_preserves_run_formatting():
    source = make_test_docx("Original summary text.", bold=True)
    decisions = {"summary": {"accepted": True, "suggested": "New summary.", "paragraph_index": 0}}

    output_path = generate(decisions, "SWE", "Google", source_path=source)
    try:
        result = Document(output_path)
        run = result.paragraphs[0].runs[0]
        assert run.text == "New summary."
        assert run.bold is True
        assert run.font.size == Pt(11)
    finally:
        os.remove(source)
        os.remove(output_path)


def test_writer_applies_accepted_experience_bullet():
    doc = Document()
    doc.add_paragraph()  # paragraph 0 — summary placeholder
    para = doc.add_paragraph("Original bullet text.")
    para.style = doc.styles["List Paragraph"] if "List Paragraph" in [s.name for s in doc.styles] else doc.styles["Normal"]

    path = tempfile.mktemp(suffix=".docx")
    doc.save(path)

    decisions = {
        "experience": [{"accepted": True, "suggested": "Improved bullet.", "paragraph_index": 1}]
    }
    output_path = generate(decisions, "SWE", "Google", source_path=path)
    try:
        result = Document(output_path)
        assert result.paragraphs[1].text == "Improved bullet."
    finally:
        os.remove(path)
        os.remove(output_path)


def test_writer_skips_rejected_experience_bullet():
    doc = Document()
    doc.add_paragraph()
    doc.add_paragraph("Original bullet text.")

    path = tempfile.mktemp(suffix=".docx")
    doc.save(path)

    decisions = {
        "experience": [{"accepted": False, "suggested": "Improved bullet.", "paragraph_index": 1}]
    }
    output_path = generate(decisions, "SWE", "Google", source_path=path)
    try:
        result = Document(output_path)
        assert result.paragraphs[1].text == "Original bullet text."
    finally:
        os.remove(path)
        os.remove(output_path)


def test_writer_reorders_skills():
    doc = Document()
    p0 = doc.add_paragraph()
    p0.add_run("Web Development: React, Node")
    p1 = doc.add_paragraph()
    p1.add_run("AI/ML: LangChain, RAG")

    path = tempfile.mktemp(suffix=".docx")
    doc.save(path)

    decisions = {
        "skills": {
            "reordered_categories": [
                {"paragraph_index": 1, "category": "AI/ML", "items": ["RAG", "LangChain"]},
                {"paragraph_index": 0, "category": "Web Development", "items": ["Node", "React"]},
            ],
            "additions": [],
        }
    }
    output_path = generate(decisions, "SWE", "Google", source_path=path)
    try:
        result = Document(output_path)
        assert result.paragraphs[0].text == "AI/ML: RAG, LangChain"
        assert result.paragraphs[1].text == "Web Development: Node, React"
    finally:
        os.remove(path)
        os.remove(output_path)


def test_writer_appends_confirmed_addition():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("AI/ML: LangChain, RAG")

    path = tempfile.mktemp(suffix=".docx")
    doc.save(path)

    decisions = {
        "skills": {
            "reordered_categories": [
                {"paragraph_index": 0, "category": "AI/ML", "items": ["LangChain", "RAG"]},
            ],
            "additions": [{"category": "AI/ML", "skill": "PyTorch", "added": True}],
        }
    }
    output_path = generate(decisions, "SWE", "Google", source_path=path)
    try:
        result = Document(output_path)
        assert "PyTorch" in result.paragraphs[0].text
    finally:
        os.remove(path)
        os.remove(output_path)


def test_writer_deletes_unselected_projects():
    doc = Document()
    doc.add_paragraph("Project A | Python")   # idx 0 — title
    doc.add_paragraph("Bullet A1")             # idx 1
    doc.add_paragraph("Project B | Java")     # idx 2 — title
    doc.add_paragraph("Bullet B1")             # idx 3

    path = tempfile.mktemp(suffix=".docx")
    doc.save(path)

    decisions = {
        "projects": {
            "all_projects": [
                {"title_paragraph_index": 0, "all_paragraph_indices": [0, 1]},
                {"title_paragraph_index": 2, "all_paragraph_indices": [2, 3]},
            ],
            "selected_title_paragraph_indices": [0],
        }
    }
    output_path = generate(decisions, "SWE", "Google", source_path=path)
    try:
        result = Document(output_path)
        texts = [p.text for p in result.paragraphs if p.text.strip()]
        assert "Project A | Python" in texts
        assert "Bullet A1" in texts
        assert "Project B | Java" not in texts
        assert "Bullet B1" not in texts
    finally:
        os.remove(path)
        os.remove(output_path)


def test_writer_keeps_selected_projects_in_original_order():
    doc = Document()
    doc.add_paragraph("Project A | Python")
    doc.add_paragraph("Bullet A1")
    doc.add_paragraph("Project B | Java")
    doc.add_paragraph("Bullet B1")
    doc.add_paragraph("Project C | Go")
    doc.add_paragraph("Bullet C1")

    path = tempfile.mktemp(suffix=".docx")
    doc.save(path)

    decisions = {
        "projects": {
            "all_projects": [
                {"title_paragraph_index": 0, "all_paragraph_indices": [0, 1]},
                {"title_paragraph_index": 2, "all_paragraph_indices": [2, 3]},
                {"title_paragraph_index": 4, "all_paragraph_indices": [4, 5]},
            ],
            "selected_title_paragraph_indices": [0, 4],
        }
    }
    output_path = generate(decisions, "SWE", "Google", source_path=path)
    try:
        result = Document(output_path)
        texts = [p.text for p in result.paragraphs if p.text.strip()]
        assert "Project A | Python" in texts
        assert "Project C | Go" in texts
        assert "Project B | Java" not in texts
        assert texts.index("Project A | Python") < texts.index("Project C | Go")
    finally:
        os.remove(path)
        os.remove(output_path)


def test_writer_drops_optional_section():
    doc = Document()
    doc.add_paragraph("RESEARCH PAPER")   # idx 0 — heading
    doc.add_paragraph("Paper citation 1") # idx 1
    doc.add_paragraph("Paper citation 2") # idx 2
    doc.add_paragraph("Other content")    # idx 3 — should stay

    path = tempfile.mktemp(suffix=".docx")
    doc.save(path)

    decisions = {
        "optional_sections": {
            "research_paper": {"keep": False, "all_paragraph_indices": [0, 1, 2]}
        }
    }
    output_path = generate(decisions, "SWE", "Google", source_path=path)
    try:
        result = Document(output_path)
        texts = [p.text for p in result.paragraphs if p.text.strip()]
        assert "RESEARCH PAPER" not in texts
        assert "Paper citation 1" not in texts
        assert "Other content" in texts
    finally:
        os.remove(path)
        os.remove(output_path)


def test_writer_keeps_optional_section():
    doc = Document()
    doc.add_paragraph("CERTIFICATIONS")
    doc.add_paragraph("NPTEL cert")

    path = tempfile.mktemp(suffix=".docx")
    doc.save(path)

    decisions = {
        "optional_sections": {
            "certifications": {"keep": True, "all_paragraph_indices": [0, 1]}
        }
    }
    output_path = generate(decisions, "SWE", "Google", source_path=path)
    try:
        result = Document(output_path)
        texts = [p.text for p in result.paragraphs if p.text.strip()]
        assert "CERTIFICATIONS" in texts
        assert "NPTEL cert" in texts
    finally:
        os.remove(path)
        os.remove(output_path)


def test_writer_projects_and_optional_deletion_uses_original_indices():
    """Both project and optional section deletions must use original indices."""
    doc = Document()
    doc.add_paragraph("Project A | Python")     # idx 0
    doc.add_paragraph("Bullet A")               # idx 1
    doc.add_paragraph("Project B | Java")       # idx 2
    doc.add_paragraph("Bullet B")               # idx 3
    doc.add_paragraph("CERTIFICATIONS")         # idx 4
    doc.add_paragraph("NPTEL cert")             # idx 5

    path = tempfile.mktemp(suffix=".docx")
    doc.save(path)

    decisions = {
        "projects": {
            "all_projects": [
                {"title_paragraph_index": 0, "all_paragraph_indices": [0, 1]},
                {"title_paragraph_index": 2, "all_paragraph_indices": [2, 3]},
            ],
            "selected_title_paragraph_indices": [0],
        },
        "optional_sections": {
            "certifications": {"keep": False, "all_paragraph_indices": [4, 5]}
        },
    }
    output_path = generate(decisions, "SWE", "Google", source_path=path)
    try:
        result = Document(output_path)
        texts = [p.text for p in result.paragraphs if p.text.strip()]
        assert "Project A | Python" in texts
        assert "Bullet A" in texts
        assert "Project B | Java" not in texts
        assert "CERTIFICATIONS" not in texts
        assert "NPTEL cert" not in texts
    finally:
        os.remove(path)
        os.remove(output_path)


def test_writer_does_not_modify_base():
    source = make_test_docx("Original summary text.")
    original_mtime = os.path.getmtime(source)
    decisions = {"summary": {"accepted": True, "suggested": "New summary.", "paragraph_index": 0}}

    output_path = generate(decisions, "SWE", "Google", source_path=source)
    try:
        assert os.path.getmtime(source) == original_mtime
        source_doc = Document(source)
        assert source_doc.paragraphs[0].text == "Original summary text."
    finally:
        os.remove(source)
        os.remove(output_path)
